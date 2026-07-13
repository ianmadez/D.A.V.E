import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_figure_placeholder(doc, label, figure_number, caption_text):
    """
    Creates a professionally formatted placeholder box for diagrams
    using a single-cell table with clean borders and light-gray background shading.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(5.5)

    # Light gray background shading
    shading_elm = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)

    # Subtle light-gray border formatting
    borders = parse_xml(
        r'<w:tcBorders {}><w:top w:val="single" w:sz="6" w:space="0" w:color="D1D5DB"/>'
        r'<w:left w:val="single" w:sz="6" w:space="0" w:color="D1D5DB"/>'
        r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="D1D5DB"/>'
        r'<w:right w:val="single" w:sz="6" w:space="0" w:color="D1D5DB"/></w:tcBorders>'
        .format(nsdecls('w'))
    )
    cell._tc.get_or_add_tcPr().append(borders)
    set_cell_margins(cell, top=280, bottom=280, left=280, right=280)

    # Internal text alignment
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)

    run = p.add_run(f"[Insert {label} Diagram Here]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Medium gray slate

    # Centered Caption directly below the box
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(12)
    caption_run = caption_p.add_run(f"Figure {figure_number}: {caption_text}")
    caption_run.font.name = 'Times New Roman'
    caption_run.font.size = Pt(10)
    caption_run.font.italic = True
    caption_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_paragraph(doc, text, space_before=0, space_after=6, is_italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = is_italic
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    run_bold = p.add_run(bold_prefix)
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(11)
    run_bold.font.bold = True

    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(11)
    return p

def add_numbered_item(doc, number, bold_prefix, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)

    run_num = p.add_run(f"{number}.  ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(11)
    run_num.font.bold = True

    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(11)
        run_bold.font.bold = True

        run_text = p.add_run(text)
        run_text.font.name = 'Times New Roman'
        run_text.font.size = Pt(11)
        return p

def add_equation_block(doc, latex_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(latex_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11.5)
    run.font.italic = True
    run.font.bold = True
    return p

def main():
    doc = Document()

    # Configure global margins (1 Inch = 2.54 cm on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # CHAPTER TITLE
        add_heading_1(doc, "CHAPTER 3: SYSTEM DESIGN AND DEVELOPMENT")

        add_paragraph(doc,
        "This chapter describes the system design, core workflow, development process, and structural models "
        "of D.A.V.E. (Direct Agentic Versioning Engine). It focuses on how the application achieves local, "
        "private, and formatting-safe codebase navigation and editing. The discussion includes the detailed "
        "software architecture, the internal mechanics of each module, the implementation of safety features, "
        "and the sequence of developmental stages."
    )

    # 3.1 SYSTEM DESIGN
    add_heading_2(doc, "3.1 System Design")

    add_paragraph(doc,
    "The architectural design of D.A.V.E. is established to prioritize privacy, processing efficiency, "
    "and safety. Operating entirely on the client machine, the system is designed to eliminate reliance "
    "on external networks while keeping resource usage within the limits of consumer-grade laptops, "
    "specifically targeting systems with 4GB of video random-access memory (VRAM)."
)

# 3.1.1 Overall D.A.V.E. Architecture
add_heading_3(doc, "3.1.1 Overall D.A.V.E. Architecture")
add_paragraph(doc, "The architecture is structured around three independent but highly coordinated modules: the Scout Module, the Plan Module, and the Execute Module.")

add_bullet(doc, "The Scout Module: ",
"Responsible for local workspace exploration, codebase tokenization, and index construction. "
"It crawls the active project workspace, creates a hierarchical file tree, and processes raw code files into discrete chunks. "
"These chunks are indexed using a local BM25 statistical retrieval engine."
)
add_bullet(doc, "The Plan Module: ",
"Acts as the cognitive engine of the system. It coordinates the tasks using a deterministic bimodal task-state system. "
"This system consists of an observable reasoning state, which holds hypothesis and logical steps, and a structured state "
"machine, which governs state transitions such as searching, analyzing, or editing."
)
add_bullet(doc, "The Execute Module: ",
"Executes safe code transformations. It manipulates the target files using Concrete Syntax Trees (CST) "
"and passes the updated files through a local verification compiler. It is supported by sanity gates and an undo manager "
"to handle automatic rollbacks if syntax anomalies or test failures are detected."
)

# 3.1.2 Frontend
add_heading_3(doc, "3.1.2 Frontend")
add_paragraph(doc, "The frontend presents a dual-entryway architecture, allowing different user archetypes to interact with the system effectively:")

add_numbered_item(doc, 1, "Command Line Interface (CLI): ",
"Designed for experienced developers who prefer a keyboard-driven, fast, and text-based interface. "
"The CLI prints real-time reasoning logs, structural file trees, and task statuses directly to the terminal using structured console outputs."
)
add_numbered_item(doc, 2, "Graphical User Interface (GUI): ",
"Constructed using the CustomTkinter toolkit. This interface presents a modern dark-mode application window, "
"visual task-state panels, dynamic file browsers, and interactive diff displays."
)
add_paragraph(doc, "To prevent graphical freeze-ups during heavy local inference workloads, the frontend communicates with the backend asynchronously using a localized thread-safe message queue.", space_before=6)

# 3.1.3 Backend
add_heading_3(doc, "3.1.3 Backend")
add_paragraph(doc, "The backend engine is written entirely in Python, running as a local daemon. It is composed of four main controllers:")

add_bullet(doc, "scanner.py: ", "Handles directory scanning, parses .gitignore rules, ignores system files, and partitions source files into tokenized blocks.")
add_bullet(doc, "planner.py: ", "Manages the agent state, handles context building, and parses unstructured LLM responses into structured execution tasks.")
add_bullet(doc, "editor.py: ", "Wraps LibCST methods to perform precise, token-safe, and formatting-safe structural edits on Python source code.")
add_bullet(doc, "terminal_runner.py: ", "Executes tests or builds commands within a localized sub-process wrapper to verify file viability post-modification.")

# 3.1.4 Local AI Models
add_heading_3(doc, "3.1.4 Local AI Models")
add_paragraph(doc, "For local inference, D.A.V.E. utilizes quantized open-source models, specifically optimized for coding and reasoning tasks. The system is designed to integrate with:")

add_bullet(doc, "Qwen-2.5-Coder (7B Parameters): ", "Quantized to 4-bit (Q4_K_M) to maintain a small memory footprint while providing competitive syntax generation.")
add_bullet(doc, "Llama-3-8B-Instruct: ", "Used as an alternative for general reasoning, text parsing, and general step planning.")
add_paragraph(doc,
"By utilizing 4-bit quantization, the model weights consume approximately 4.3GB of system memory. "
"When running active inference, the active context is constrained using a dynamic sliding context window, "
"keeping the VRAM usage under the target threshold of 4GB.",
space_before=6
)

# 3.1.5 File System
add_heading_3(doc, "3.1.5 File System")
add_paragraph(doc,
"D.A.V.E. does not require a complex relational database engine, which reduces installation friction and makes the system highly portable. "
"All metadata, indexes, and session details are written directly to a hidden directory named .dave_cache within the root folder of the user's active project. "
"The cache folder contains:"
)
add_bullet(doc, "index.json: ", "Stores BM25 vocabulary, term document frequencies, and relative file system paths.")
add_bullet(doc, "history.json: ", "Holds a chronological stack of changes, including compressed pre-edit backups of altered files, enabling multi-step rollbacks.")
add_bullet(doc, "metadata.json: ", "Preserves workspace configuration, excluded directories, and model selection preferences.")

# 3.1.6 LLM Providers
add_heading_3(doc, "3.1.6 LLM Providers")
add_paragraph(doc, "Inference is facilitated locally using the Ollama platform. D.A.V.E. interacts with Ollama via an internal client class that targets the local host port (usually http://localhost:11434/api/generate). This request handler includes:")

add_bullet(doc, "Adaptive Backoff Logic: ", "Automatically retries queries up to five times with progressive delays (1s, 2s, 4s, 8s, 16s) to handle hardware congestion.")
add_bullet(doc, "Strict JSON Formatting: ", "Enforces structured outputs from the LLM by passing JSON schemas in the system instruction prompts, ensuring that the response matches the expected internal types of the parser.")

# 3.1.7 Security and Local-First Design
add_heading_3(doc, "3.1.7 Security and Local-First Design")
add_paragraph(doc,
"Because D.A.V.E. processes all computations locally, the source code, system environment variables, "
"database schemas, and intellectual property remain entirely on the user's physical machine. "
"This design successfully eliminates common cloud security risks, such as data leaks, server compromises, "
"and API credential theft. Furthermore, the local executor is restricted from executing arbitrary system "
"commands unless they are explicitly declared inside the test verification rules by the user."
)

# 3.2 SYSTEM WORKFLOW
add_heading_2(doc, "3.2 System Workflow")
add_paragraph(doc, "The operations of D.A.V.E. follow a structured and predictable workflow. Unlike unstructured autonomous agents that rely on open-ended loops, D.A.V.E. divides its execution stages into precise, verifiable steps.")

# 3.2.1 User Workflow
add_heading_3(doc, "3.2.1 User Workflow")
add_paragraph(doc, "The user interaction follows a straightforward progression:")
add_numbered_item(doc, 1, "Workspace Selection: ", "The user initializes D.A.V.E. and selects a target project folder.")
add_numbered_item(doc, 2, "Indexing: ", "The system performs an initial scan and indexing of all source files, storing the index in .dave_cache.")
add_numbered_item(doc, 3, "Task Specification: ", "The user enters a natural language command (for example: \"Refactor the database connection to use a singleton pattern\").")
add_numbered_item(doc, 4, "Plan Review: ", "The user reviews the step-by-step plan generated by the Plan Module, showing which files are slated for modifications.")
add_numbered_item(doc, 5, "Execution and Validation: ", "Upon receiving user approval, the system applies the modifications, runs the validation checks, and displays the outcome.")
add_numbered_item(doc, 6, "Review and Commit: ", "The user reviews the visual diff. If the change is unsatisfactory or fails tests, the user can click \"Undo\" to restore the previous state.")

# 3.2.2 Agent Workflow
add_heading_3(doc, "3.2.2 Agent Workflow")
add_paragraph(doc,
"Internally, the agent operates through an iterative loop of scanning, reasoning, editing, and validating. "
"When a task is accepted, the system shifts into the Scout State. It executes a keyword search using BM25, "
"pulling the most relevant code chunks into the active memory.\n\n"
"Next, the system transitions to the Plan State. The local model analyzes the retrieved code, generates a hypothesis, "
"and builds a list of sequential changes.\n\n"
"Once the plan is generated, the agent moves into the Execute State. For each step, it generates structural code edits, "
"performs syntax verification, writes the safe changes via LibCST, and executes terminal checks. If any step fails, "
"the workflow pauses, logs the error, and provides the option to revert."
)

# 3.2.3 Request Processing
add_heading_3(doc, "3.2.3 Request Processing")
add_paragraph(doc, "When a user issues a request, the agent handles the context window carefully to prevent memory exhaustion:")

add_numbered_item(doc, 1, "Retrieval Phase: ",
"The user query is tokenized, and the BM25 engine extracts the top scoring blocks. "
"The BM25 score is computed using the following formula:"
)

# Mathematical Equation Block
add_equation_block(doc,
"Score(D, Q) = Sum_{i=1}^{n} [ IDF(q_i) * ( f(q_i, D) * (k_1 + 1) ) / ( f(q_i, D) + k_1 * (1 - b + b * ( |D| / avgdl )) ) ]"
)

add_paragraph(doc,
"where f(q_i, D) is the term frequency of query token q_i in code chunk D, |D| is the length of the chunk in words, "
"avgdl is the average length of all chunks in the codebase, and k_1 and b are structural parameters set to 1.5 and 0.75 respectively.",
space_after=8
)
add_numbered_item(doc, 2, "Context Assembly: ", "The highest-ranked chunks are formatted with XML tags and prepended with system rules, creating a consolidated prompt.")
add_numbered_item(doc, 3, "Model Inference: ", "The prompt is sent to Ollama. The response is parsed to extract the files to be read and the exact operations to be conducted.")

# 3.2.4 File Editing Process
add_heading_3(doc, "3.2.4 File Editing Process")
add_paragraph(doc, "To modify source code safely, D.A.V.E. implements formatting-safe code surgery. Standard text generation or regex replacements are highly prone to syntax breakage and formatting loss. D.A.V.E. solves this through the following process:")
add_numbered_item(doc, 1, "", "The target file is loaded, and LibCST parses the raw string into a Concrete Syntax Tree.")
add_numbered_item(doc, 2, "", "A specialized subclass of libcst.CSTTransformer is instantiated with the instructions generated by the Plan Module.")
add_numbered_item(doc, 3, "", "The transformer navigates the nodes of the tree, locating specific function definitions, class declarations, or variables.")
add_numbered_item(doc, 4, "", "The transformer performs node replacements, additions, or deletions on the tree structure.")
add_numbered_item(doc, 5, "", "The modified tree is serialized back into Python source code.")
add_paragraph(doc,
"Because the library manipulates syntax nodes directly rather than raw text, the surrounding comments, "
"line breaks, indentation levels, and coding styles are completely preserved, keeping the codebase clean "
"and structurally valid.",
space_before=6
)

# 3.3 DEVELOPMENT PROCESS
add_heading_2(doc, "3.3 Development Process")
add_paragraph(doc, "The development of D.A.V.E. followed an iterative, agile methodology. Each iteration focused on refining local execution safety, optimizing memory usage, and building a highly cohesive user interface.")

# 3.3.1 Planning
add_heading_3(doc, "3.3.1 Planning")
add_paragraph(doc,
"The initial phase involved reviewing existing frameworks such as Cursor and OpenClaw. "
"The investigation identified a distinct research gap: the lack of local-first, low-overhead, "
"and formatting-safe editing options for developers with restricted hardware. Based on these findings, "
"the core requirements were defined, specifying that the tool must operate completely offline, "
"keep the RAM usage of the LLM under 4GB, and use a robust syntax manipulator to prevent syntax corruption."
)

# 3.3.2 UI Design
add_heading_3(doc, "3.3.2 UI Design")
add_paragraph(doc, "The interface development was divided into two distinct tracks:")
add_bullet(doc, "The CLI layout ", "was designed first, ensuring that all backend actions could be monitored through clean log messages and visual markers.")
add_bullet(doc, "The GUI layout ", "was designed with a multi-pane approach. It features a left-hand navigation pane for codebase directory trees, a center panel for chat logs and planning actions, and a right-hand panel dedicated to the live task list, reasoning log, and validation output.")
add_paragraph(doc, "Modern design tokens, such as soft borders, contrasting text fields, and clear success/error highlights, were implemented using CustomTkinter styles.", space_before=6)

# 3.3.3 Backend Development
add_heading_3(doc, "3.3.3 Backend Development")
add_paragraph(doc, "The backend modules were developed incrementally:")
add_numbered_item(doc, 1, "Scanner Component: ", "Developed first to ensure accurate directory traversal, file size filtering, and clean exclusion of binary assets and dependency folders.")
add_numbered_item(doc, 2, "Indexing Module: ", "Developed around the rank_bm25 library, utilizing standard tokenizers to index Python files.")
add_numbered_item(doc, 3, "Editor Module: ", "Focused on wrapping LibCST code transformers. Custom visitors were constructed to handle safe insertions of functions, modifications of assignment statements, and updates to import declarations.")

# 3.3.4 AI Integration
add_heading_3(doc, "3.3.4 AI Integration")
add_paragraph(doc,
"Integrating local LLMs through Ollama required comprehensive prompt engineering. "
"Standard general-purpose prompts often caused smaller models to hallucinate file paths or output invalid JSON structures. "
"This challenge was resolved by introducing strict system directives, instructing the model to act as a state-aware "
"agent and format its plans in a highly structured structure. Furthermore, sliding context windows were implemented "
"to prune stale chat history, which prevented memory errors on systems with 4GB VRAM."
)

# 3.3.5 Testing
add_heading_3(doc, "3.3.5 Testing")
add_paragraph(doc, "Comprehensive unit, integration, and performance tests were executed:")
add_bullet(doc, "Unit Testing: ", "Focused on validating the tokenization of files, verifying correct BM25 score calculations, and testing the LibCST transformer classes on diverse Python scripts.")
add_bullet(doc, "Integration Testing: ", "Evaluated the interaction between the Plan Module, the Execute Module, and the terminal runner. The system was tested against scenarios where invalid code was injected, validating that the pre-execution sanity gate successfully aborted the write operations.")
add_bullet(doc, "Hardware Benchmarking: ", "Profiled VRAM and system memory utilization during active indexing and heavy model inference.")

# 3.3.6 Improvements
add_heading_3(doc, "3.3.6 Improvements")
add_paragraph(doc,
"Several enhancements were made based on the testing outcomes. "
"First, an automatic fallback was added to catch and parse corrupted JSON output from the LLM, "
"ensuring the planning state would not crash. Second, to handle files with irregular character encodings, "
"the scanning engine was updated with robust encoding detection, falling back gracefully to UTF-8 or ISO-8859-1. "
"Finally, an idempotency guard was implemented to limit the number of successive tool calls, preventing the agent "
"from entering infinite execution loops when attempting to fix failing tests."
)

# 3.4 LIST OF FIGURES
add_heading_2(doc, "3.4 List of Figures")
add_paragraph(doc, "This section presents the structural models, workflow behaviors, and technical diagrams of the D.A.V.E. system.")

# 3.4.1 System Flowchart
add_heading_3(doc, "3.4.1 System Flowchart")
add_paragraph(doc, "The system flowchart illustrates the high-level logic of the application. It maps the path a user request takes from initial submission to final verification, showing the decision loops and safety gates that govern system execution.")
add_figure_placeholder(doc, "System Flowchart", "3.1", "D.A.V.E. System Flowchart Placeholder.")
add_paragraph(doc,
"The flowchart shown in Figure 3.1 tracks the entire processing lifecycle of a user instruction. "
"The process begins when the user submits a natural language request through the interface. "
"The system analyzes the workspace files and builds a local index.\n\n"
"Once indexed, the search engine ranks the files to locate the relevant target functions or files. "
"The planning module then defines the sequence of edits and presents them as a structured plan.\n\n"
"If the user approves the plan, the executor applies the edits to the target files using LibCST. "
"The system then runs a post-execution sanity check. If the code is structurally sound and compiles, "
"the edits are finalized, and the user is notified. If an error is detected, the rollback mechanism is triggered, "
"restoring the codebase to its original state."
)

# 3.4.2 Entity Relationship Diagram (ERD)
add_heading_3(doc, "3.4.2 Entity Relationship Diagram (ERD)")
add_paragraph(doc, "The Entity Relationship Diagram (ERD) defines the logical database structure of D.A.V.E., outlining how local projects, sessions, files, prompts, and history logs are associated.")
add_figure_placeholder(doc, "ERD", "3.2", "D.A.V.E. Entity Relationship Diagram Placeholder.")
add_paragraph(doc,
"The ERD in Figure 3.2 shows the data relationships maintained in the .dave_cache directory. "
"The primary entity is the Project, which represents the root directory of the active workspace. "
"A Project is associated with multiple Session records, tracking each time a user starts the application.\n\n"
"Each Session contains multiple Codebase Files, which are parsed and indexed. "
"Each file is associated with specific Term Vectors used by the BM25 search engine.\n\n"
"Furthermore, each Session tracks user Prompts and their resulting Task Plans. "
"A Task Plan contains multiple planned steps, which map to individual Edit History records. "
"The Edit History entity stores pre-edit file hashes and compressed backup paths, "
"allowing each modification to be reversed to its previous state."
)

# 3.4.3 Class Diagram
add_heading_3(doc, "3.4.3 Class Diagram")
add_paragraph(doc, "The class diagram displays the object-oriented structure of the D.A.V.E. system, demonstrating the relationships, attributes, and methods of the core software components.")
add_figure_placeholder(doc, "Class", "3.3", "D.A.V.E. Class Diagram Placeholder.")
add_paragraph(doc,
"The class diagram in Figure 3.3 represents the internal architecture of the backend controllers. "
"The primary controller is the Agent class, which aggregates the UIController, Planner, FileManager, Editor, "
"and LLMProvider classes. The FileManager handles directory crawling, parsing, and caching, "
"using the IndexEngine helper class to compute BM25 scores.\n\n"
"The Planner maintains the system state, processing natural language requests into structured lists of tasks. "
"The Editor uses specialized subclasses of libcst.CSTTransformer to perform code modifications.\n\n"
"The LLMProvider connects with the local Ollama daemon to handle text generation, and the UIController coordinates "
"display updates, translating agent state changes into visual progress logs on the CustomTkinter GUI."
)

# 3.4.4 Activity Diagram
add_heading_3(doc, "3.4.4 Activity Diagram")
add_paragraph(doc, "The activity diagram maps the lifecycle of a user request, depicting the concurrent actions and conditions that occur as the system analyzes, edits, and validates source code.")
add_figure_placeholder(doc, "Activity", "3.4", "D.A.V.E. Activity Diagram Placeholder.")
add_paragraph(doc,
"The activity diagram in Figure 3.4 illustrates the parallel and sequential activities executed when a request is processed. "
"Once a user submits a prompt, the system concurrently queries the BM25 search engine and parses the active workspace structure. "
"The agent then evaluates the context size. If the text size exceeds the local model limitations, "
"the system triggers sliding window pruning to compress the history.\n\n"
"The planning module then generates the structural task plan. The system encounters a decision node requiring user confirmation. "
"If the user rejects the proposed steps, the activity returns to the planning phase or terminates.\n\n"
"Upon approval, the executor applies changes to the syntax tree. The terminal runner executes verification scripts. "
"If the tests succeed, the task transitions to a completed state. If they fail, the rollback path is executed, "
"and the error details are sent to the log panel."
)

# 3.4.5 Sequence Diagram
add_heading_3(doc, "3.4.5 Sequence Diagram")
add_paragraph(doc, "The sequence diagram illustrates the chronological communication flow across the system components, tracing the execution path from the user's initial interaction to the file system output.")
add_figure_placeholder(doc, "Sequence", "3.5", "D.A.V.E. Sequence Diagram Placeholder.")
add_paragraph(doc,
"The sequence diagram in Figure 3.5 displays the time-sequenced interactions of a codebase update task. "
"The interaction begins with the User entering a request on the GUI. The GUI passes this payload to the Agent controller. "
"The Agent requests relevant files from the FileManager, which queries the IndexEngine and returns the code blocks.\n\n"
"The Agent then forwards these blocks along with the prompt to the LLMProvider. "
"The LLMProvider returns a structured execution plan.\n\n"
"The Agent presents this plan to the GUI for the User to review. "
"Once the User confirms, the Agent commands the Editor to modify the code. "
"The Editor parses the target file, applies the transformer edits, and writes the modifications back to the file system.\n\n"
"Finally, the Agent calls the TerminalRunner to execute the tests, and the final results are returned to the GUI for the User to review."
)

# 3.4.6 Use Case Diagram
add_heading_3(doc, "3.4.6 Use Case Diagram")
add_paragraph(doc, "The use case diagram outlines the functional scope of D.A.V.E., showing how the user acts as the primary actor to trigger core system operations.")
add_figure_placeholder(doc, "Use Case", "3.6", "D.A.V.E. Use Case Diagram Placeholder.")
add_paragraph(doc,
"The use case diagram in Figure 3.6 defines the functional capabilities available to the user. "
"The primary actor is the Developer, who can initiate several key actions. "
"The core use cases include Importing a Project Workspace, which automatically includes the sub-use case of Indexing Codebase Files.\n\n"
"The user can Inspect Codebase Structure to locate functions or files, as well as Edit Source Files. "
"The editing use case includes the mandatory validation step of Verifying Syntax Safety, "
"and optionally includes Executing Local Test Suites.\n\n"
"Additional use cases allow the user to Revert Code Changes, View Edit History, and Configure Local LLM Models, "
"giving them complete, local control over the software tool's execution behavior."
)

# Save Document
output_filename = "chapter_3_system_design.docx"
doc.save(output_filename)
print(f"Success! Generated a professional Word document at: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    main()